import configparser
import json
import sys
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from fontTools.ttLib import TTFont
from docx import Document

if getattr(sys, "frozen", False):
    BASE_DIR = Path(sys.executable).parent
else:
    BASE_DIR = Path(__file__).parent
CONFIG_FILE = BASE_DIR / "config.ini"

_DEFAULTS = {
    "dpi": "300",
    "font_size": "12",
    "darkness_threshold": "200",
    "baseline_font": "Arial",
    "fonts_dir": r"C:\Windows\Fonts",
    "line_spacing_factor": "1.15",
    "sample_text": "sample_text.txt",
    "fonts_list": "fonts.txt",
    "output_dir": "output",
}


def ensure_config():
    if CONFIG_FILE.exists():
        return
    cp = configparser.ConfigParser()
    cp["Settings"] = _DEFAULTS
    with open(CONFIG_FILE, "w") as f:
        cp.write(f)


def load_config():
    ensure_config()
    cp = configparser.ConfigParser()
    cp.read(CONFIG_FILE)
    s = cp["Settings"]
    dpi = int(s.get("dpi", _DEFAULTS["dpi"]))
    font_size = int(s.get("font_size", _DEFAULTS["font_size"]))
    margin = int(1 * dpi)
    output_dir = BASE_DIR / s.get("output_dir", _DEFAULTS["output_dir"])
    return {
        "dpi": dpi,
        "font_size": font_size,
        "darkness_threshold": int(s.get("darkness_threshold", _DEFAULTS["darkness_threshold"])),
        "baseline_font": s.get("baseline_font", _DEFAULTS["baseline_font"]),
        "fonts_dir": Path(s.get("fonts_dir", _DEFAULTS["fonts_dir"])),
        "line_spacing_factor": float(s.get("line_spacing_factor", _DEFAULTS["line_spacing_factor"])),
        "sample_text": BASE_DIR / s.get("sample_text", _DEFAULTS["sample_text"]),
        "fonts_list": BASE_DIR / s.get("fonts_list", _DEFAULTS["fonts_list"]),
        "output_dir": output_dir,
        "results_json": output_dir / "results.json",
        "results_docx": output_dir / "results.docx",
        "page_width": int(8.5 * dpi),
        "page_height": int(11 * dpi),
        "margin": margin,
        "text_width": int(8.5 * dpi) - 2 * margin,
        "font_size_px": int(font_size * dpi / 72),
    }


def load_font_list(cfg):
    path = cfg["fonts_list"]
    if not path.exists():
        return []
    names = set()
    for line in path.read_text(encoding="utf-8").splitlines():
        name = line.strip()
        if name:
            names.add(name)
    return sorted(names)


def find_fonts(cfg, names):
    target = set(names)
    found = {}
    for p in sorted(Path(cfg["fonts_dir"]).iterdir()):
        if p.suffix.lower() not in (".ttf", ".otf", ".ttc"):
            continue
        for idx in range(16):
            try:
                family = ImageFont.truetype(str(p), 20, index=idx).getname()[0]
            except Exception:
                break
            if family not in target or family in found:
                continue
            try:
                tt = TTFont(str(p), fontNumber=idx)
                os2 = tt.get("OS/2")
                tt.close()
                if os2 and not (350 <= os2.usWeightClass <= 450):
                    continue
            except Exception:
                continue
            found[family] = (p, idx)
    return found


def _wrap_text(text, font, text_width):
    lines = []
    for paragraph in text.split("\n"):
        if not paragraph.strip():
            lines.append("")
            continue
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            test = current + " " + word
            bbox = font.getbbox(test)
            if bbox[2] - bbox[0] <= text_width:
                current = test
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def render_and_measure(text, font_path, font_index, cfg):
    fspx = cfg["font_size_px"]
    font = ImageFont.truetype(str(font_path), fspx, index=font_index)
    line_height = int(fspx * cfg["line_spacing_factor"])
    wrapped = _wrap_text(text, font, cfg["text_width"])
    usable = cfg["page_height"] - 2 * cfg["margin"]
    lines_per_page = usable // line_height

    dark = 0
    for start in range(0, len(wrapped), lines_per_page):
        chunk = wrapped[start : start + lines_per_page]
        img = Image.new("L", (cfg["page_width"], cfg["page_height"]), 255)
        draw = ImageDraw.Draw(img)
        y = cfg["margin"]
        for line in chunk:
            draw.text((cfg["margin"], y), line, font=font, fill=0)
            y += line_height
        dark += sum(img.histogram()[:cfg["darkness_threshold"]])
    return dark


def process_font(text, font_path, font_index, cfg):
    try:
        return render_and_measure(text, font_path, font_index, cfg)
    except Exception:
        return None


def compute_relative(results, baseline):
    base_ink = results.get(baseline)
    if not base_ink:
        return {}
    return {name: round(ink / base_ink * 100, 1) for name, ink in results.items()}


def save_json(cfg, results, relative):
    data = {}
    for name in sorted(results):
        data[name] = {
            "dark_pixels": results[name],
            "ink_vs_baseline": relative.get(name, "—"),
        }
    cfg["output_dir"].mkdir(exist_ok=True)
    cfg["results_json"].write_text(json.dumps(data, indent=2), encoding="utf-8")


def save_docx(cfg, results, relative):
    doc = Document()
    doc.add_heading("Ink Ranker Results", level=1)
    table = doc.add_table(rows=1, cols=3, style="Light List Accent 1")
    hdr = table.rows[0].cells
    hdr[0].text = "Font"
    hdr[1].text = "Dark Pixels"
    hdr[2].text = "vs Baseline (%)"
    for name, ink in sorted(results.items(), key=lambda x: x[1]):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"{ink:,}"
        row[2].text = str(relative.get(name, "—"))
    cfg["output_dir"].mkdir(exist_ok=True)
    doc.save(str(cfg["results_docx"]))
