import json
import sys
import time
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from fontTools.ttLib import TTFont
from docx import Document
from settings import (
    DPI, FONT_SIZE, DARKNESS_THRESHOLD,
    BASELINE_FONT, FONTS_DIR, LINE_SPACING_FACTOR,
)

if getattr(sys, "frozen", False):
    BUNDLE_DIR = Path(getattr(sys, '_MEIPASS', '.'))
    BASE_DIR = Path(sys.executable).parent
else:
    BUNDLE_DIR = Path(__file__).parent
    BASE_DIR = Path(__file__).parent

FONTS_DIR = Path(FONTS_DIR)


def load_config():
    margin = DPI
    font_size_px = int(FONT_SIZE * DPI / 72)
    return {
        "dpi": DPI,
        "font_size_px": font_size_px,
        "darkness_threshold": DARKNESS_THRESHOLD,
        "baseline_font": BASELINE_FONT,
        "fonts_dir": FONTS_DIR,
        "line_spacing_factor": LINE_SPACING_FACTOR,
        "sample_text": BUNDLE_DIR / "sample.txt",
        "fonts_list": BUNDLE_DIR / "fonts.txt",
        "results_json": BASE_DIR / "results.json",
        "results_docx": BASE_DIR / "results.docx",
        "page_width": int(8.5 * DPI),
        "page_height": int(11 * DPI),
        "margin": margin,
        "text_width": int(8.5 * DPI) - 2 * margin,
    }


def load_font_list(cfg):
    path = cfg["fonts_list"]
    if not path.exists():
        return []
    names = set()
    for line in path.read_text(encoding="utf-8").splitlines():
        name = line.strip()
        if name:
            names.add(name)
    return sorted(names)


def find_fonts(cfg, names):
    target = set(names)
    found = {}
    for p in sorted(Path(cfg["fonts_dir"]).iterdir()):
        if p.suffix.lower() not in (".ttf", ".otf", ".ttc"):
            continue
        for idx in range(16):
            try:
                family = ImageFont.truetype(str(p), 20, index=idx).getname()[0]
            except Exception:
                break
            if family not in target or family in found:
                continue
            try:
                tt = TTFont(str(p), fontNumber=idx)
                os2 = tt.get("OS/2")
                tt.close()
                if os2 and not (350 <= os2.usWeightClass <= 450):
                    continue
            except Exception:
                continue
            found[family] = (p, idx)
        time.sleep(0)
        if len(found) == len(target):
            break
    return found


def _wrap_text(text, font, text_width):
    lines = []
    for paragraph in text.split("\n"):
        if not paragraph.strip():
            lines.append("")
            continue
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            test = current + " " + word
            bbox = font.getbbox(test)
            if bbox[2] - bbox[0] <= text_width:
                current = test
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def render_and_measure(text, font_path, font_index, cfg):
    fspx = cfg["font_size_px"]
    font = ImageFont.truetype(str(font_path), fspx, index=font_index)
    line_height = int(fspx * cfg["line_spacing_factor"])
    wrapped = _wrap_text(text, font, cfg["text_width"])
    usable = cfg["page_height"] - 2 * cfg["margin"]
    lines_per_page = usable // line_height

    dark = 0
    for start in range(0, len(wrapped), lines_per_page):
        chunk = wrapped[start : start + lines_per_page]
        img = Image.new("L", (cfg["page_width"], cfg["page_height"]), 255)
        draw = ImageDraw.Draw(img)
        y = cfg["margin"]
        for line in chunk:
            draw.text((cfg["margin"], y), line, font=font, fill=0)
            y += line_height
        dark += sum(img.histogram()[:cfg["darkness_threshold"]])
    return dark


def process_font(text, font_path, font_index, cfg):
    try:
        return render_and_measure(text, font_path, font_index, cfg)
    except Exception:
        return None


def compute_relative(results, baseline):
    base_ink = results.get(baseline)
    if not base_ink:
        return {}
    return {name: round(ink / base_ink * 100, 1) for name, ink in results.items()}


def save_json(cfg, results, relative, scan_seconds=0):
    data = {"scan_time": f"{scan_seconds}s"}
    for name in sorted(results):
        data[name] = {
            "dark_pixels": results[name],
            "ink_vs_baseline": relative.get(name, "—"),
        }

    cfg["results_json"].write_text(json.dumps(data, indent=2), encoding="utf-8")


def save_docx(cfg, results, relative):
    doc = Document()
    doc.add_heading("Ink Ranker Results", level=1)
    table = doc.add_table(rows=1, cols=3, style="Light List Accent 1")
    hdr = table.rows[0].cells
    hdr[0].text = "Font"
    hdr[1].text = "Dark Pixels"
    hdr[2].text = "vs Baseline (%)"
    for name, ink in sorted(results.items(), key=lambda x: x[1]):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"{ink:,}"
        row[2].text = str(relative.get(name, "—"))

    doc.save(str(cfg["results_docx"]))
